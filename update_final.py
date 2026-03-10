import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil

# Copy the pristine template (or the user created file) to final output
shutil.copy('vzor.docx', 'cv05_koma_modular_final.docx')

doc = docx.Document('cv05_koma_modular_final.docx')

def set_cell_answers(cell, heading_text, answer_lines):
    # keep the first paragraph (heading)
    p0 = cell.paragraphs[0]
    # delete the rest
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    
    # add new paragraphs
    for line in answer_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')
        rPr.append(rFonts)

# TABLE 31: LinkedIn Text
linkedin_text = """Modul isn't a container.
A hotel v Letňanech u PVA EXPO to dokazuje lépe než jakýkoli argument.

Většina developerů, se kterými mluvím, má v hlavě stejný obrázek:
šedá buňka, plech, dočasnost.
Zatím jejich konkurent položil první modul v prvním červencovém týdnu.
O necelé tři měsíce později dolaďoval interiéry.

Co konkrétně vzniklo:
→ 4podlažní hotel, 104 pokojů, restaurace
→ 160 modulů vyrobených ve Vizovicích
→ hlučná montáž jeřábem: jen několik dní
→ pevná cena, žádné vícepráce
→ sítě vedené z chodby — servis bez vstupu do pokojů

Hosté neví, že spí v modulární stavbě.
Architekti to poznají. A oceňují to.

Budoucnost výstavby není otázka materiálu.
Je to otázka přístupu.
Jak dlouho trvala vaše poslední stavba oproti původnímu plánu?
Napište do komentáře. 👇

#modularita #development #hotely #výstavba #KOMAmodular""".split('\n')

set_cell_answers(doc.tables[31].cell(0,0), "", linkedin_text)

# TABLE 34: IG Format
ig_format = doc.tables[34].cell(1,1)
for p in ig_format.paragraphs:
    p._element.getparent().remove(p._element)
p = ig_format.add_paragraph()
r = p.add_run("Jednoduchý obrázkový příspěvek (Single Image)")
r.font.name = 'Arial'
r.font.size = Pt(11)

# TABLE 35: IG Format Reason
ig_format_reason = [
    "Instagram feed post (single image) s důrazem na vizuál.",
    "Formát maximalizuje šanci na okamžité sdílení (Share), což je pro aktuální algoritmus Instagramu klíčové."
]
set_cell_answers(doc.tables[35].cell(0,0), "", ig_format_reason)

# TABLE 36: IG Visual Concept
ig_visual = [
    "Jednoduchý post s nahraným obrázkem (vizuál prémiové modulární stavby).",
    "Fotografie funguje jako hlavní důkaz ('Show, don\\'t tell')."
]
set_cell_answers(doc.tables[36].cell(0,0), "", ig_visual)

# TABLE 37: IG Text
ig_text = """Stavební buňka? ❌ 
Takhle vypadá modulární architektura dnes.

⏱️ 4 měsíce výstavby.
💰 Pevná cena.
☔️ Žádné čekání na počasí.

Znáš developera, který pořád staví postaru a ztrácí čas?
✈️ Pošli mu to do zpráv.""".split('\n')
set_cell_answers(doc.tables[37].cell(0,0), "", ig_text)

doc.save('cv05_koma_modular_final.docx')
print("Successfully generated cv05_koma_modular_final.docx!")
